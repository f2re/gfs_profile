from __future__ import annotations

import tempfile
import unittest
from dataclasses import dataclass, field
from datetime import datetime, timedelta, timezone
from pathlib import Path
from unittest.mock import patch

import numpy as np
from PIL import Image
from docx import Document

from meteogram_report import (
    MeteogramReportError,
    build_meteogram_report_data,
    write_meteogram_report,
)


@dataclass
class _Source:
    source_id: str = "gefs"
    label: str = "GEFS 0.25° · NOAA/NCEP"
    model: str = "NOAA GEFS 0.25°"
    provider: str = "NOAA/NCEP через Open-Meteo"
    resolution: str = "0.25°"
    ensemble: bool = True


@dataclass
class _Series:
    times: list[datetime]
    fields: dict[str, np.ndarray]
    stats: dict[str, dict[str, np.ndarray]]
    daily_dates: list
    daily_stats: dict[str, dict[str, np.ndarray]]
    source: _Source = field(default_factory=_Source)
    point_label: str = "Санкт-Петербург"
    requested_lat: float = 59.9391
    requested_lon: float = 30.3159
    grid_lat: float = 60.0
    grid_lon: float = 30.25
    timezone: str = "Europe/Moscow"
    retrieved_at_utc: datetime = field(default_factory=lambda: datetime.now(timezone.utc))
    member_count: int = 31
    expected_member_count: int = 31
    warnings: list[str] = field(default_factory=list)
    sampling_mode: str = "raw_model_grid"

    def values(self, name: str) -> np.ndarray:
        return self.fields.get(name, np.full(len(self.times), np.nan))

    def statistic(self, name: str, stat: str) -> np.ndarray:
        return self.stats.get(name, {}).get(stat, np.full(len(self.times), np.nan))

    def daily_statistic(self, name: str, stat: str) -> np.ndarray:
        return self.daily_stats.get(name, {}).get(stat, np.full(len(self.daily_dates), np.nan))


def _series(hours: int = 120) -> _Series:
    start = datetime(2026, 8, 14, 0, tzinfo=timezone(timedelta(hours=3)))
    times = [start + timedelta(hours=3 * index) for index in range(hours // 3 + 1)]
    count = len(times)
    phase = np.linspace(0, 5 * np.pi, count)
    temperature = 14 + 7 * np.sin(phase)
    precipitation = np.zeros(count)
    precipitation[10:13] = [0.3, 1.2, 0.5]
    fields = {
        "temperature_2m": temperature,
        "dew_point_2m": temperature - 5,
        "relative_humidity_2m": np.clip(70 - 20 * np.sin(phase), 35, 98),
        "precipitation": precipitation,
        "wind_speed_10m": np.full(count, 5.0),
        "wind_gusts_10m": np.full(count, 9.0),
        "wind_direction_10m": np.full(count, 240.0),
        "pressure_msl": np.linspace(1018, 1008, count),
        "cloud_cover": np.linspace(35, 95, count),
        "weather_code": np.where(precipitation > 0, 61, 3),
        "ensemble_member_count": np.full(count, 31.0),
        "precipitation_accumulation_hours": np.full(count, 3.0),
        "precipitation_probability_0p1": np.where(precipitation > 0, 70.0, 15.0),
        "precipitation_probability_1": np.where(precipitation >= 1.0, 45.0, 5.0),
        "precipitation_probability_5": np.zeros(count),
    }
    stats = {
        "temperature_2m": {"q10": temperature - 3, "q90": temperature + 3},
        "wind_gusts_10m": {"q90": np.full(count, 13.0)},
    }
    dates = list(dict.fromkeys(item.date() for item in times))
    q50 = np.linspace(0.0, 4.0, len(dates))
    daily_stats = {
        "precipitation": {
            "q10": np.maximum(0, q50 - 0.5),
            "q50": q50,
            "q90": q50 + 2.0,
            "coverage_hours": np.full(len(dates), 24.0),
            "complete_day": np.ones(len(dates)),
        }
    }
    return _Series(times, fields, stats, dates, daily_stats)


class MeteogramReportTests(unittest.TestCase):
    def test_report_data_has_daily_and_control_tables(self) -> None:
        data = build_meteogram_report_data(_series())
        self.assertIn("Ансамблевый", data.title)
        self.assertGreaterEqual(len(data.daily_rows), 5)
        self.assertGreaterEqual(len(data.control_rows), 13)
        text = " ".join(
            [*data.main_lines]
            + [row.precipitation for row in data.daily_rows]
            + [row.ensemble for row in data.daily_rows]
        )
        self.assertIn("q10-q90", text)
        self.assertIn("член", text)
        self.assertIn("мм/3 ч", text)
        self.assertRegex(text, r"\(\d+/31\)")
        self.assertNotIn("nan", text.lower())

    def test_docx_contains_narrative_chart_and_tables(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            directory = Path(tmp)
            chart = directory / "meteogram.png"
            Image.new("RGB", (1400, 760), "white").save(chart)
            result = write_meteogram_report(_series(), chart, "docx", output_dir=directory)
            self.assertTrue(result.path.is_file())
            document = Document(result.path)
            self.assertGreaterEqual(len(document.tables), 3)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("Главное", text)
            self.assertIn("Метеограмма", text)
            self.assertIn("Методика и ограничения", text)

    def test_pdf_request_works_without_libreoffice(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            directory = Path(tmp)
            chart = directory / "meteogram.png"
            Image.new("RGB", (1400, 760), "white").save(chart)
            with patch("meteogram_report._find_libreoffice", return_value=None):
                result = write_meteogram_report(
                    _series(),
                    chart,
                    "pdf",
                    output_dir=directory,
                    pdf_fallback_to_docx=False,
                )
            self.assertEqual(result.format, "pdf")
            self.assertIsNone(result.fallback_reason)
            self.assertTrue(result.path.is_file())
            self.assertGreater(result.path.stat().st_size, 1500)
            self.assertEqual(result.path.read_bytes()[:5], b"%PDF-")

    def test_telegram_output_parser_and_keyboard(self) -> None:
        import telegram_meteogram as module

        output, cleaned = module._extract_output_format(
            "Санкт-Петербург ensemble=gefs days=5 format=pdf"
        )
        self.assertEqual(output, "pdf")
        self.assertNotIn("format=", cleaned)
        callbacks = [
            button.callback_data
            for row in module._output_keyboard().inline_keyboard
            for button in row
        ]
        self.assertIn("meteo:format:png", callbacks)
        self.assertIn("meteo:format:docx", callbacks)
        self.assertIn("meteo:format:pdf", callbacks)
        with self.assertRaises(Exception):
            module._extract_output_format("Москва format=pdf output=docx")


if __name__ == "__main__":
    unittest.main()
