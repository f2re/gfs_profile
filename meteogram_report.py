from __future__ import annotations

"""DOCX/PDF report generation for a single meteogram series.

The module deliberately depends only on the public ``MeteogramSeries`` shape:
``source``, ``times``, ``values()``, ``statistic()`` and optional daily
statistics.  It does not fetch data and never combines different ensemble
systems.  DOCX and PDF are rendered from the same report data. PDF is rendered
natively with Matplotlib; LibreOffice is retained only as an optional fallback.
"""

import math
import os
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field
from datetime import date, datetime, timezone
from pathlib import Path
from typing import Any, Iterable, Sequence

import numpy as np

REPORT_FORMATS = frozenset({"docx", "pdf"})
REPORT_FONT = os.getenv("METEOGRAM_REPORT_FONT", "Liberation Sans")
PDF_TIMEOUT_SECONDS = max(10, int(os.getenv("METEOGRAM_PDF_TIMEOUT", "90")))

THUNDER_CODES = frozenset({95, 96, 99})
FOG_CODES = frozenset({45, 48})
DRIZZLE_CODES = frozenset({51, 53, 55, 56, 57})
FREEZING_CODES = frozenset({56, 57, 66, 67})
RAIN_CODES = frozenset({61, 63, 65, 66, 67, 80, 81, 82})
SNOW_CODES = frozenset({71, 73, 75, 77, 85, 86})

WEEKDAYS_RU = ("пн", "вт", "ср", "чт", "пт", "сб", "вс")
WIND_DIRECTIONS_RU = (
    "С",
    "ССВ",
    "СВ",
    "ВСВ",
    "В",
    "ВЮВ",
    "ЮВ",
    "ЮЮВ",
    "Ю",
    "ЮЮЗ",
    "ЮЗ",
    "ЗЮЗ",
    "З",
    "ЗСЗ",
    "СЗ",
    "ССЗ",
)


class MeteogramReportError(RuntimeError):
    """Report creation or conversion failed."""


@dataclass(frozen=True, slots=True)
class ReportDay:
    day: date
    weather: str
    temperature: str
    precipitation: str
    wind: str
    pressure: str
    ensemble: str


@dataclass(frozen=True, slots=True)
class ReportControlTime:
    time: datetime
    temperature: str
    humidity_cloud: str
    precipitation: str
    wind: str
    pressure: str
    ensemble: str


@dataclass(frozen=True, slots=True)
class MeteogramReportData:
    title: str
    subtitle: str
    point_line: str
    period_line: str
    source_line: str
    main_lines: tuple[str, ...]
    daily_rows: tuple[ReportDay, ...]
    control_rows: tuple[ReportControlTime, ...]
    method_lines: tuple[str, ...]
    warning_lines: tuple[str, ...]
    filename_stem: str


@dataclass(frozen=True, slots=True)
class MeteogramReportResult:
    path: Path
    format: str
    docx_path: Path
    fallback_reason: str | None = None
    cleanup_paths: tuple[Path, ...] = field(default_factory=tuple)


def normalise_report_format(value: str | None) -> str:
    key = str(value or "docx").strip().lower().lstrip(".")
    aliases = {
        "word": "docx",
        "document": "docx",
        "документ": "docx",
        "ворд": "docx",
        "portable": "pdf",
    }
    key = aliases.get(key, key)
    if key not in REPORT_FORMATS:
        raise MeteogramReportError("Формат отчёта должен быть DOCX или PDF")
    return key


def build_meteogram_report_data(series: Any) -> MeteogramReportData:
    if not getattr(series, "times", None):
        raise MeteogramReportError("Нет прогностических сроков для отчёта")

    times = list(series.times)
    if len(times) < 2:
        raise MeteogramReportError("Для отчёта требуется минимум два прогностических срока")

    source = series.source
    ensemble = bool(getattr(source, "ensemble", False))
    kind = "Ансамблевый модельный прогноз" if ensemble else "Модельный прогноз"
    point_label = _clean_text(getattr(series, "point_label", "Точка прогноза"))
    title = f"{kind}: {point_label}"
    subtitle = str(getattr(source, "model", getattr(source, "label", "Модель")))

    requested_lat = _float_or_none(getattr(series, "requested_lat", None))
    requested_lon = _float_or_none(getattr(series, "requested_lon", None))
    grid_lat = _float_or_none(getattr(series, "grid_lat", None))
    grid_lon = _float_or_none(getattr(series, "grid_lon", None))
    point_bits = [point_label]
    if requested_lat is not None and requested_lon is not None:
        point_bits.append(f"запрошено {requested_lat:.4f}, {requested_lon:.4f}")
    if grid_lat is not None and grid_lon is not None:
        point_bits.append(f"расчётная точка {grid_lat:.3f}, {grid_lon:.3f}")
    point_line = " · ".join(point_bits)

    timezone_name = str(getattr(series, "timezone", "местное время"))
    period_line = (
        f"Период: {times[0]:%d.%m.%Y %H:%M} - {times[-1]:%d.%m.%Y %H:%M}"
        f" · местное время ({timezone_name})"
    )

    provider = str(getattr(source, "provider", "источник не указан"))
    resolution = getattr(source, "resolution", None)
    retrieved = getattr(series, "retrieved_at_utc", None)
    source_bits = [provider]
    if resolution:
        source_bits.append(str(resolution))
    if isinstance(retrieved, datetime):
        source_bits.append(f"получено {retrieved.astimezone(timezone.utc):%d.%m.%Y %H:%M} UTC")
    source_line = " · ".join(source_bits)

    daily_rows = tuple(_build_daily_rows(series))
    control_rows = tuple(_build_control_rows(series))
    main_lines = tuple(_build_main_lines(series, daily_rows))
    method_lines = tuple(_build_method_lines(series))
    warning_lines = tuple(_build_warning_lines(series))

    start = times[0]
    model_key = _safe_filename(str(getattr(source, "source_id", getattr(source, "model", "model"))))
    point_key = _safe_filename(point_label)
    filename_stem = f"Прогноз_{point_key}_{model_key}_{start:%Y-%m-%d}"

    return MeteogramReportData(
        title=title,
        subtitle=subtitle,
        point_line=point_line,
        period_line=period_line,
        source_line=source_line,
        main_lines=main_lines,
        daily_rows=daily_rows,
        control_rows=control_rows,
        method_lines=method_lines,
        warning_lines=warning_lines,
        filename_stem=filename_stem,
    )


def write_meteogram_report(
    series: Any,
    chart_path: str | Path,
    output_format: str,
    *,
    output_dir: str | Path | None = None,
    pdf_fallback_to_docx: bool = True,
) -> MeteogramReportResult:
    """Create a DOCX or PDF report and return all temporary paths for cleanup.

    DOCX and PDF use the same ``MeteogramReportData``. PDF is rendered
    natively with Matplotlib, so normal operation does not require LibreOffice.
    LibreOffice remains a secondary converter if the native renderer fails.
    """

    fmt = normalise_report_format(output_format)
    data = build_meteogram_report_data(series)
    chart = Path(chart_path)
    if not chart.is_file() or chart.stat().st_size <= 0:
        raise MeteogramReportError("PNG метеограммы не найден или пуст")

    out_dir = Path(output_dir) if output_dir is not None else Path(tempfile.mkdtemp(prefix="gfs_meteogram_report_"))
    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / f"{data.filename_stem}.docx"
    write_meteogram_docx(series, chart, docx_path, report_data=data)

    cleanup = [docx_path]
    if fmt == "docx":
        return MeteogramReportResult(
            path=docx_path,
            format="docx",
            docx_path=docx_path,
            cleanup_paths=tuple(cleanup),
        )

    pdf_path = out_dir / f"{data.filename_stem}.pdf"
    native_error: str | None = None
    try:
        from meteogram_pdf import MeteogramPdfError, write_meteogram_pdf

        write_meteogram_pdf(data, chart, pdf_path)
    except (MeteogramPdfError, OSError, ValueError) as exc:
        native_error = str(exc)
        try:
            convert_docx_to_pdf(docx_path, pdf_path)
        except MeteogramReportError as libreoffice_exc:
            combined = (
                f"Нативный PDF: {native_error}; "
                f"резервный LibreOffice: {libreoffice_exc}"
            )
            if not pdf_fallback_to_docx:
                raise MeteogramReportError(combined) from libreoffice_exc
            return MeteogramReportResult(
                path=docx_path,
                format="docx",
                docx_path=docx_path,
                fallback_reason=combined,
                cleanup_paths=tuple(cleanup),
            )

    cleanup.append(pdf_path)
    return MeteogramReportResult(
        path=pdf_path,
        format="pdf",
        docx_path=docx_path,
        cleanup_paths=tuple(cleanup),
    )


def write_meteogram_docx(
    series: Any,
    chart_path: str | Path,
    output_path: str | Path,
    *,
    report_data: MeteogramReportData | None = None,
) -> Path:
    """Write a landscape A4 report with narrative, tables and full meteogram."""

    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Mm, Pt
    except ImportError as exc:  # pragma: no cover - exercised on deployment without dependency
        raise MeteogramReportError(
            "Для DOCX установите зависимость python-docx>=1.1,<2"
        ) from exc

    data = report_data or build_meteogram_report_data(series)
    chart = Path(chart_path)
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)
    section.header_distance = Mm(4)
    section.footer_distance = Mm(4)

    normal = document.styles["Normal"]
    normal.font.name = REPORT_FONT
    normal.font.size = Pt(8.4)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0
    _set_style_font(normal, REPORT_FONT)

    for style_name, size, bold, color in (
        ("Title", 17, True, "17324D"),
        ("Heading 1", 12, True, "17324D"),
        ("Heading 2", 10, True, "17324D"),
    ):
        style = document.styles[style_name]
        style.font.name = REPORT_FONT
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = _rgb(color)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(3)
        _set_style_font(style, REPORT_FONT)

    if "Meteo Small" not in document.styles:
        small = document.styles.add_style("Meteo Small", WD_STYLE_TYPE.PARAGRAPH)
    else:
        small = document.styles["Meteo Small"]
    small.font.name = REPORT_FONT
    small.font.size = Pt(7.2)
    small.font.color.rgb = _rgb("52616B")
    small.paragraph_format.space_after = Pt(1)
    _set_style_font(small, REPORT_FONT)

    core = document.core_properties
    core.title = data.title
    core.subject = "Модельный метеорологический прогноз"
    core.author = "GFS Profile Telegram bot"
    core.keywords = "метеограмма, ансамбль, модельный прогноз, GFS Profile"
    core.comments = "Документ сформирован автоматически; данные являются модельным прогнозом."

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("GFS Profile · модельный прогноз · не наблюдение")
    run.font.name = REPORT_FONT
    run.font.size = Pt(7)
    run.font.color.rgb = _rgb("60717B")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("GFS Profile · ")
    footer_run.font.name = REPORT_FONT
    footer_run.font.size = Pt(7)
    _append_field(footer, "PAGE")

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run(data.title)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(2)
    subtitle_run = subtitle.add_run(data.subtitle)
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(10.5)
    subtitle_run.font.color.rgb = _rgb("214A63")

    for line in (data.point_line, data.period_line, data.source_line):
        document.add_paragraph(line, style="Meteo Small")

    disclaimer = document.add_paragraph()
    disclaimer.paragraph_format.space_before = Pt(2)
    disclaimer.paragraph_format.space_after = Pt(4)
    disclaimer_run = disclaimer.add_run(
        "Важно: один модельный ансамбль или одна модель; не радиозонд, не станция и не официальный выпуск предупреждения."
    )
    disclaimer_run.bold = True
    disclaimer_run.font.size = Pt(8)
    disclaimer_run.font.color.rgb = _rgb("8A3E25")

    document.add_heading("Главное", level=1)
    summary_table = document.add_table(rows=1, cols=1)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_table.autofit = True
    summary_cell = summary_table.cell(0, 0)
    summary_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_shading(summary_cell, "EEF5F8")
    _set_cell_margins(summary_cell, top=90, start=130, bottom=90, end=130)
    for index, line in enumerate(data.main_lines or ("Существенные явления не выделены.",)):
        paragraph = summary_cell.paragraphs[0] if index == 0 else summary_cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.add_run("• " + line)

    document.add_heading("Прогноз по местным суткам", level=1)
    daily_headers = (
        "Дата",
        "Характер погоды",
        "Температура",
        "Осадки",
        "Ветер",
        "Давление",
        "Ансамбль",
    )
    daily_widths = (18, 37, 39, 58, 43, 31, 41)
    daily_table = document.add_table(rows=1, cols=len(daily_headers))
    daily_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    daily_table.style = "Table Grid"
    _configure_table(daily_table, daily_headers, daily_widths)
    for row in data.daily_rows:
        cells = daily_table.add_row().cells
        values = (
            f"{row.day:%d.%m}\n{WEEKDAYS_RU[row.day.weekday()]}",
            row.weather,
            row.temperature,
            row.precipitation,
            row.wind,
            row.pressure,
            row.ensemble,
        )
        _fill_table_row(cells, values, font_size=7.15)

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading("Метеограмма", level=1)
    caption = document.add_paragraph(
        "Центральная линия показывает среднее или медиану согласно параметру; полосы ансамбля - q25-q75 и q10-q90. "
        "Осадки не сглаживаются. Время на оси местное."
    )
    caption.style = document.styles["Meteo Small"]
    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.space_before = Pt(2)
    picture_paragraph.paragraph_format.space_after = Pt(0)
    try:
        picture_paragraph.add_run().add_picture(str(chart), width=Mm(268))
    except Exception as exc:
        raise MeteogramReportError(f"Не удалось вставить PNG метеограммы в DOCX: {exc}") from exc

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading("Прогноз по контрольным срокам", level=1)
    control_note = document.add_paragraph(
        "Шаг: через 6 часов в первые 72 часа, далее через 12 часов. Значения относятся к ближайшему доступному сроку модели."
    )
    control_note.style = document.styles["Meteo Small"]
    control_headers = (
        "Дата и время",
        "T / Td",
        "RH / облачность",
        "Осадки",
        "Ветер",
        "Давление",
        "Ансамбль",
    )
    control_widths = (28, 34, 38, 61, 50, 28, 36)
    control_table = document.add_table(rows=1, cols=len(control_headers))
    control_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    control_table.style = "Table Grid"
    _configure_table(control_table, control_headers, control_widths)
    for row in data.control_rows:
        cells = control_table.add_row().cells
        values = (
            f"{row.time:%d.%m}\n{row.time:%H:%M}",
            row.temperature,
            row.humidity_cloud,
            row.precipitation,
            row.wind,
            row.pressure,
            row.ensemble,
        )
        _fill_table_row(cells, values, font_size=6.85)

    document.add_heading("Методика и ограничения", level=1)
    for line in data.method_lines:
        paragraph = document.add_paragraph(style="Meteo Small")
        paragraph.add_run("• " + line)
    if data.warning_lines:
        document.add_heading("Предупреждения о данных", level=2)
        for line in data.warning_lines:
            paragraph = document.add_paragraph(style="Meteo Small")
            run = paragraph.add_run("• " + line)
            run.font.color.rgb = _rgb("8A3E25")

    # Keep the report compact while retaining a readable fallback when the
    # table spans an extra page in LibreOffice.
    for paragraph in document.paragraphs:
        paragraph.paragraph_format.widow_control = True

    temporary = output.with_suffix(output.suffix + ".tmp")
    document.save(temporary)
    if not temporary.is_file() or temporary.stat().st_size <= 0:
        raise MeteogramReportError("python-docx не создал непустой документ")
    temporary.replace(output)
    return output


def convert_docx_to_pdf(
    docx_path: str | Path,
    output_path: str | Path,
    *,
    timeout_seconds: int | None = None,
) -> Path:
    source = Path(docx_path)
    target = Path(output_path)
    if not source.is_file() or source.stat().st_size <= 0:
        raise MeteogramReportError("DOCX для преобразования в PDF отсутствует")

    binary = _find_libreoffice()
    if binary is None:
        raise MeteogramReportError(
            "PDF недоступен: LibreOffice Writer не установлен; сформирован DOCX"
        )

    target.parent.mkdir(parents=True, exist_ok=True)
    timeout = timeout_seconds or PDF_TIMEOUT_SECONDS
    with tempfile.TemporaryDirectory(prefix="gfs_lo_profile_") as profile_dir:
        profile_uri = Path(profile_dir).resolve().as_uri()
        env = os.environ.copy()
        env.setdefault("SAL_USE_VCLPLUGIN", "svp")
        env["HOME"] = profile_dir
        command = [
            binary,
            "--headless",
            "--nologo",
            "--nodefault",
            "--nolockcheck",
            f"-env:UserInstallation={profile_uri}",
            "--convert-to",
            "pdf:writer_pdf_Export",
            "--outdir",
            str(target.parent),
            str(source),
        ]
        try:
            result = subprocess.run(
                command,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=timeout,
                check=False,
                env=env,
            )
        except subprocess.TimeoutExpired as exc:
            raise MeteogramReportError(
                f"PDF не сформирован за {timeout} с; сформирован DOCX"
            ) from exc
        except OSError as exc:
            raise MeteogramReportError(
                f"Не удалось запустить LibreOffice: {exc}; сформирован DOCX"
            ) from exc

    produced = target.parent / f"{source.stem}.pdf"
    if result.returncode != 0 or not produced.is_file() or produced.stat().st_size <= 0:
        details = (result.stderr or result.stdout or "неизвестная ошибка").strip()
        if len(details) > 350:
            details = details[-350:]
        raise MeteogramReportError(
            f"LibreOffice не сформировал PDF: {details}; сформирован DOCX"
        )
    if produced.resolve() != target.resolve():
        if target.exists():
            target.unlink()
        produced.replace(target)
    return target


def _build_daily_rows(series: Any) -> list[ReportDay]:
    times = list(series.times)
    grouped: dict[date, list[int]] = {}
    for index, current in enumerate(times):
        grouped.setdefault(current.date(), []).append(index)

    temperature = _values(series, "temperature_2m")
    dewpoint = _values(series, "dew_point_2m")
    precipitation = _values(series, "precipitation")
    wind = _values(series, "wind_speed_10m")
    gust = _values(series, "wind_gusts_10m")
    pressure = _values(series, "pressure_msl")
    humidity = _values(series, "relative_humidity_2m")
    cloud = _cloud_values(series)
    weather_code = _values(series, "weather_code")
    member_counts = _values(series, "ensemble_member_count")

    ensemble = bool(getattr(series.source, "ensemble", False))
    expected = int(getattr(series, "expected_member_count", 0) or getattr(series, "member_count", 0) or 0)
    result: list[ReportDay] = []
    for day_index, (current_day, indices) in enumerate(grouped.items()):
        t_low = _nanmin(temperature[indices])
        t_high = _nanmax(temperature[indices])
        temperature_text = _format_range(t_low, t_high, "°C", signed=True)
        if ensemble:
            q10 = _statistic(series, "temperature_2m", "q10")
            q90 = _statistic(series, "temperature_2m", "q90")
            spread_low = _nanmin(q10[indices])
            spread_high = _nanmax(q90[indices])
            if spread_low is not None and spread_high is not None:
                temperature_text += f"\nq10-q90 {_format_range(spread_low, spread_high, '°C', signed=True)}"

        precip_q50 = _daily_stat(series, "precipitation", "q50", day_index)
        precip_q10 = _daily_stat(series, "precipitation", "q10", day_index)
        precip_q90 = _daily_stat(series, "precipitation", "q90", day_index)
        coverage = _daily_stat(series, "precipitation", "coverage_hours", day_index)
        complete = _daily_stat(series, "precipitation", "complete_day", day_index)
        period_label = "сут" if complete is not None and bool(complete) else _coverage_period(coverage)
        has_daily_members = precip_q50 is not None
        if precip_q50 is None:
            precip_q50 = _nansum(precipitation[indices])
        if complete is None:
            fallback_coverage = _fallback_day_coverage_hours(series, indices)
            period_label = (
                "сут"
                if fallback_coverage is not None and 22.5 <= fallback_coverage <= 25.5
                else _coverage_period(fallback_coverage)
            )
        precipitation_text = _format_precipitation_amount(precip_q50, period_label)
        if ensemble and precip_q10 is not None and precip_q90 is not None:
            precipitation_text += "\n" + _format_precipitation_spread(
                precip_q10, precip_q90, period_label
            )
        elif ensemble and not has_daily_members:
            precipitation_text += "\nсумма центрального ряда"
        probability_parts = _probability_parts(series, indices)
        if probability_parts:
            precipitation_text += "\n" + "\n".join(probability_parts)

        max_wind = _nanmax(wind[indices])
        max_gust = _nanmax(gust[indices])
        if ensemble:
            gust_q90 = _statistic(series, "wind_gusts_10m", "q90")
            q90_max = _nanmax(gust_q90[indices])
        else:
            q90_max = None
        wind_text = _format_wind(max_wind, max_gust, q90_max)

        pressure_text = _format_range(
            _nanmin(pressure[indices]), _nanmax(pressure[indices]), "гПа"
        )

        min_members = _nanmin(member_counts[indices]) if ensemble else None
        if ensemble:
            observed = int(round(min_members)) if min_members is not None else int(getattr(series, "member_count", 0) or 0)
            member_text = f"{observed}/{expected or observed} членов"
            signal_probability = _max_probability(series, indices)
            member_text += f"\n{_signal_label(signal_probability)}"
        else:
            member_text = "одна модель"

        weather_text = _weather_label(
            ensemble=ensemble,
            weather_codes=weather_code[indices],
            cloud_values=cloud[indices],
            humidity_values=humidity[indices],
            precipitation_total=float(precip_q50 or 0.0),
            precipitation_upper=precip_q90,
            max_probability=_max_probability(series, indices),
        )
        result.append(
            ReportDay(
                day=current_day,
                weather=weather_text,
                temperature=temperature_text,
                precipitation=precipitation_text,
                wind=wind_text,
                pressure=pressure_text,
                ensemble=member_text,
            )
        )
    return result


def _build_control_rows(series: Any) -> list[ReportControlTime]:
    times = list(series.times)
    selected = _control_indices(times)
    temperature = _values(series, "temperature_2m")
    dewpoint = _values(series, "dew_point_2m")
    humidity = _values(series, "relative_humidity_2m")
    cloud = _cloud_values(series)
    precipitation = _values(series, "precipitation")
    wind = _values(series, "wind_speed_10m")
    gust = _values(series, "wind_gusts_10m")
    direction = _values(series, "wind_direction_10m")
    pressure = _values(series, "pressure_msl")
    member_counts = _values(series, "ensemble_member_count")
    intervals = _values(series, "precipitation_accumulation_hours")
    ensemble = bool(getattr(series.source, "ensemble", False))
    expected = int(getattr(series, "expected_member_count", 0) or getattr(series, "member_count", 0) or 0)

    result: list[ReportControlTime] = []
    for index in selected:
        t = _finite_at(temperature, index)
        td = _finite_at(dewpoint, index)
        temp_text = _fmt_signed(t, 1) + " °C" if t is not None else "—"
        if td is not None:
            temp_text += f"\nTd {_fmt_signed(td, 1)} °C"

        rh = _finite_at(humidity, index)
        cl = _finite_at(cloud, index)
        humidity_text = _fmt(rh, 0) + " %" if rh is not None else "—"
        if cl is not None:
            humidity_text += f"\nоблачность {_fmt(cl, 0)} %"

        precip = _finite_at(precipitation, index)
        interval = _finite_at(intervals, index)
        interval_label = f"/{_fmt(interval, 0)} ч" if interval and interval > 0 else "/интервал"
        precipitation_text = (
            f"{_fmt(precip, 1)} мм{interval_label}" if precip is not None and precip >= 0.05 else "без существенных осадков"
        )
        probability_parts = _probability_parts(series, [index])
        if probability_parts:
            precipitation_text += "\n" + "\n".join(probability_parts)

        speed = _finite_at(wind, index)
        gust_value = _finite_at(gust, index)
        direction_value = _finite_at(direction, index)
        direction_label = _wind_direction_label(direction_value)
        wind_text = (
            f"{direction_label}, {_fmt(speed, 1)} м/с" if speed is not None else "—"
        )
        if gust_value is not None:
            wind_text += f"\nпорывы {_fmt(gust_value, 1)} м/с"
        if ensemble:
            q90 = _finite_at(_statistic(series, "wind_gusts_10m", "q90"), index)
            if q90 is not None:
                wind_text += f"\nq90 {_fmt(q90, 1)} м/с"

        pressure_value = _finite_at(pressure, index)
        pressure_text = f"{_fmt(pressure_value, 0)} гПа" if pressure_value is not None else "—"

        if ensemble:
            count = _finite_at(member_counts, index)
            observed = int(round(count)) if count is not None else int(getattr(series, "member_count", 0) or 0)
            member_text = f"{observed}/{expected or observed}"
        else:
            member_text = "1 модель"

        result.append(
            ReportControlTime(
                time=times[index],
                temperature=temp_text,
                humidity_cloud=humidity_text,
                precipitation=precipitation_text,
                wind=wind_text,
                pressure=pressure_text,
                ensemble=member_text,
            )
        )
    return result


def _build_main_lines(series: Any, daily_rows: Sequence[ReportDay]) -> list[str]:
    times = list(series.times)
    temperature = _values(series, "temperature_2m")
    wind = _values(series, "wind_speed_10m")
    gust = _values(series, "wind_gusts_10m")
    humidity = _values(series, "relative_humidity_2m")
    ensemble = bool(getattr(series.source, "ensemble", False))

    lines: list[str] = []
    min_index = _nanargmin(temperature)
    max_index = _nanargmax(temperature)
    if min_index is not None and max_index is not None:
        line = (
            f"Температура центрального ряда: {_fmt_signed(temperature[min_index], 1)} °C "
            f"({times[min_index]:%d.%m %H:%M}) - {_fmt_signed(temperature[max_index], 1)} °C "
            f"({times[max_index]:%d.%m %H:%M})."
        )
        if ensemble:
            q10 = _statistic(series, "temperature_2m", "q10")
            q90 = _statistic(series, "temperature_2m", "q90")
            low = _nanmin(q10)
            high = _nanmax(q90)
            if low is not None and high is not None:
                line += f" Диапазон q10-q90 по срокам: {_fmt_signed(low, 1)}…{_fmt_signed(high, 1)} °C."
        lines.append(line)

    daily_amounts: list[tuple[float, str]] = []
    for row in daily_rows:
        match = re.match(r"([0-9]+(?:,[0-9]+)?)", row.precipitation)
        if match:
            daily_amounts.append((float(match.group(1).replace(",", ".")), f"{row.day:%d.%m}"))
    if daily_amounts:
        amount, day_label = max(daily_amounts)
        daily_kind = (
            "суточная медиана осадков"
            if ensemble and _has_daily_statistic(series, "precipitation", "q50")
            else "сумма осадков центрального ряда"
        )
        line = f"Максимальная {daily_kind}: {_fmt(amount, 1)} мм {day_label}."
    else:
        line = "Существенная суточная сумма осадков центрального ряда не выделяется."
    if ensemble:
        max_probability = _max_probability(series, range(len(times)))
        if max_probability is not None:
            line += f" Максимальный сигнал осадков по членам: {_fmt(max_probability, 0)} % ({_signal_label(max_probability).lower()})."
    lines.append(line)

    max_wind_index = _nanargmax(wind)
    risk_gust = _statistic(series, "wind_gusts_10m", "q90") if ensemble else gust
    max_gust_index = _nanargmax(risk_gust)
    wind_parts = []
    if max_wind_index is not None:
        wind_parts.append(
            f"средний ветер до {_fmt(wind[max_wind_index], 1)} м/с {times[max_wind_index]:%d.%m %H:%M}"
        )
    if max_gust_index is not None:
        prefix = "q90 порывов" if ensemble else "порывы"
        wind_parts.append(
            f"{prefix} до {_fmt(risk_gust[max_gust_index], 1)} м/с {times[max_gust_index]:%d.%m %H:%M}"
        )
    if wind_parts:
        lines.append("Ветер: " + "; ".join(wind_parts) + ".")

    high_humidity = np.flatnonzero(np.isfinite(humidity) & (humidity >= 95.0))
    if high_humidity.size:
        first = int(high_humidity[0])
        lines.append(
            f"Относительная влажность достигает 95 % или выше начиная с {times[first]:%d.%m %H:%M}; это модельный индикатор, а не наблюдение тумана."
        )

    if ensemble:
        counts = _values(series, "ensemble_member_count")
        minimum = _nanmin(counts)
        observed = int(round(minimum)) if minimum is not None else int(getattr(series, "member_count", 0) or 0)
        expected = int(getattr(series, "expected_member_count", 0) or observed)
        lines.append(
            f"Полнота ансамбля по срокам: не менее {observed}/{expected or observed} членов. "
            "Оценка описывает разброс внутри одной ансамблевой системы и не является межмодельным консенсусом."
        )
    return lines[:6]


def _build_method_lines(series: Any) -> list[str]:
    ensemble = bool(getattr(series.source, "ensemble", False))
    lines = [
        "Все сроки и суточные границы приведены к местному времени точки.",
        "Осадки показаны за исходный интервал и не сглаживаются; графические тренды не повышают разрешение модели.",
        "Направление ветра метеорологическое - откуда дует.",
        "Источник и расчётная точка указаны в заголовке; профиль не является наблюдением или радиозондом.",
    ]
    if ensemble:
        lines.insert(
            1,
            "Центр ансамбля: среднее для T/Td/давления и медиана для остальных параметров; полосы - q25-q75 и q10-q90.",
        )
        lines.append(
            "Вероятности осадков - доля доступных членов выбранного ансамбля, превысивших порог за исходный интервал."
        )
        if not _has_daily_statistic(series, "precipitation", "q50"):
            lines.append(
                "Поставщик не передал member-by-member суточную статистику; в таблице дана сумма центрального ряда, явно отмеченная в ячейке."
            )
    sampling = str(getattr(series, "sampling_mode", ""))
    if sampling == "raw_model_grid":
        lines.append("Использована модельная ячейка; высотная коррекция поставщика отключена.")
    return lines


def _build_warning_lines(series: Any) -> list[str]:
    result = []
    for item in list(getattr(series, "warnings", ()) or ()):
        text = _clean_text(item)
        if text and text not in result:
            result.append(text)
    return result[:8]


def _weather_label(
    *,
    ensemble: bool,
    weather_codes: np.ndarray,
    cloud_values: np.ndarray,
    humidity_values: np.ndarray,
    precipitation_total: float,
    precipitation_upper: float | None,
    max_probability: float | None,
) -> str:
    codes = {
        int(round(value))
        for value in np.asarray(weather_codes, dtype=float)
        if np.isfinite(value)
    }
    if not ensemble:
        if codes & THUNDER_CODES:
            return "Гроза / ливневые осадки"
        if codes & FREEZING_CODES:
            return "Переохлаждённые осадки"
        if codes & SNOW_CODES:
            return "Снег / снежные заряды"
        if codes & RAIN_CODES:
            return "Дождь / ливневые осадки"
        if codes & DRIZZLE_CODES:
            return "Морось"
        if codes & FOG_CODES:
            return "Туман по коду модели"

    wet_upper = precipitation_upper if precipitation_upper is not None else precipitation_total
    if ensemble and max_probability is not None:
        if max_probability >= 70.0 and wet_upper >= 0.1:
            return "Устойчивый сигнал осадков"
        if max_probability >= 40.0 and wet_upper >= 0.1:
            return "Осадки вероятны"
        if max_probability >= 15.0 and wet_upper >= 0.1:
            return "Отдельные сценарии осадков"

    cloud = _nanmean(cloud_values)
    humidity = _nanmean(humidity_values)
    if humidity is not None and humidity >= 95.0 and (cloud or 0.0) >= 70.0:
        return "Очень влажно, сплошная облачность"
    if cloud is None:
        return "Характер погоды не определён"
    if cloud < 20.0:
        return "Малооблачно"
    if cloud < 60.0:
        return "Переменная облачность"
    if cloud < 85.0:
        return "Облачно"
    return "Пасмурно"


def _probability_parts(series: Any, indices: Iterable[int]) -> list[str]:
    index_list = list(indices)
    result = []
    intervals = _values(series, "precipitation_accumulation_hours")
    member_counts = _values(series, "ensemble_member_count")
    default_members = int(getattr(series, "member_count", 0) or 0)
    for field, label in (
        ("precipitation_probability_0p1", "P≥0,1 мм"),
        ("precipitation_probability_1", "P≥1 мм"),
        ("precipitation_probability_5", "P≥5 мм"),
    ):
        values = _values(series, field)
        finite_indices = [
            index
            for index in index_list
            if 0 <= index < len(values) and np.isfinite(values[index])
        ]
        if not finite_indices:
            continue
        best_index = max(finite_indices, key=lambda index: float(values[index]))
        maximum = float(values[best_index])
        interval = _finite_at(intervals, best_index)
        interval_label = (
            f"/{_fmt(interval, 0)} ч"
            if interval is not None and interval > 0
            else "/интервал"
        )
        members = _finite_at(member_counts, best_index)
        denominator = int(round(members)) if members is not None else default_members
        count_label = ""
        if denominator > 0:
            events = int(round(maximum * denominator / 100.0))
            count_label = f" ({events}/{denominator})"
        result.append(
            f"{label}{interval_label}: {_fmt(maximum, 0)} %{count_label}"
        )
    return result


def _max_probability(series: Any, indices: Iterable[int]) -> float | None:
    index_list = list(indices)
    values = _values(series, "precipitation_probability_0p1")
    return _nanmax(values[index_list]) if index_list else None


def _signal_label(probability: float | None) -> str:
    if probability is None:
        return "сигнал не рассчитан"
    if probability >= 70.0:
        return "устойчивый сигнал"
    if probability >= 40.0:
        return "вероятный сигнал"
    if probability >= 15.0:
        return "отдельные сценарии"
    return "слабый сигнал"


def _control_indices(times: Sequence[datetime]) -> list[int]:
    if not times:
        return []
    seconds = np.asarray([item.timestamp() for item in times], dtype=float)
    start = seconds[0]
    end_hours = max(0.0, (seconds[-1] - start) / 3600.0)
    targets: list[float] = []
    hour = 0.0
    while hour <= min(72.0, end_hours) + 0.01:
        targets.append(hour)
        hour += 6.0
    hour = 84.0
    while hour <= end_hours + 0.01:
        targets.append(hour)
        hour += 12.0
    indices: list[int] = []
    for target in targets:
        index = int(np.argmin(np.abs((seconds - start) / 3600.0 - target)))
        if not indices or index != indices[-1]:
            indices.append(index)
    if indices and indices[-1] != len(times) - 1 and end_hours - targets[-1] >= 6.0:
        indices.append(len(times) - 1)
    return indices


def _format_wind(max_wind: float | None, max_gust: float | None, q90_gust: float | None) -> str:
    parts = []
    if max_wind is not None:
        parts.append(f"до {_fmt(max_wind, 1)} м/с")
    if max_gust is not None:
        parts.append(f"порывы {_fmt(max_gust, 1)} м/с")
    if q90_gust is not None:
        parts.append(f"q90 {_fmt(q90_gust, 1)} м/с")
    return "\n".join(parts) if parts else "—"


def _format_precipitation_amount(value: float | None, period: str) -> str:
    if value is None or not math.isfinite(value) or value < 0.05:
        return "без существенных осадков"
    return f"{_fmt(value, 1)} {_precipitation_period_suffix(period)}"


def _format_precipitation_spread(low: float, high: float, period: str) -> str:
    return (
        f"q10-q90 {_fmt(low, 1)}…{_fmt(high, 1)} "
        f"{_precipitation_period_suffix(period)}"
    )


def _precipitation_period_suffix(period: str) -> str:
    clean = str(period or "").strip()
    if clean == "сут":
        return "мм/сут"
    if clean.endswith("ч"):
        return f"мм за {clean}"
    return "мм за доступный период"


def _coverage_period(value: float | None) -> str:
    if value is None or not math.isfinite(value) or value <= 0:
        return "доступный период"
    return f"{_fmt(value, 0)} ч"


def _fallback_day_coverage_hours(series: Any, indices: Sequence[int]) -> float | None:
    if not indices:
        return None
    intervals = _values(series, "precipitation_accumulation_hours")
    selected = np.asarray(
        [intervals[index] for index in indices if 0 <= index < len(intervals)],
        dtype=float,
    )
    selected = selected[np.isfinite(selected) & (selected > 0)]
    if selected.size:
        return float(np.sum(selected))
    times = list(getattr(series, "times", ()))
    if len(times) < 2:
        return None
    steps = np.asarray(
        [
            (right.timestamp() - left.timestamp()) / 3600.0
            for left, right in zip(times, times[1:])
        ],
        dtype=float,
    )
    steps = steps[np.isfinite(steps) & (steps > 0)]
    if not steps.size:
        return None
    return float(len(indices) * np.median(steps))


def _has_daily_statistic(series: Any, parameter: str, statistic: str) -> bool:
    method = getattr(series, "daily_statistic", None)
    values = None
    if callable(method):
        try:
            values = method(parameter, statistic)
        except Exception:
            values = None
    if values is None:
        values = (
            getattr(series, "daily_stats", {})
            .get(parameter, {})
            .get(statistic, [])
        )
    array = np.asarray(values, dtype=float)
    return bool(array.size and np.isfinite(array).any())


def _daily_stat(series: Any, parameter: str, statistic: str, index: int) -> float | None:
    method = getattr(series, "daily_statistic", None)
    values = None
    if callable(method):
        try:
            values = method(parameter, statistic)
        except Exception:
            values = None
    if values is None:
        values = (
            getattr(series, "daily_stats", {})
            .get(parameter, {})
            .get(statistic, np.asarray([], dtype=float))
        )
    array = np.asarray(values, dtype=float)
    return _finite_at(array, index)


def _values(series: Any, name: str) -> np.ndarray:
    try:
        values = series.values(name)
    except Exception:
        values = getattr(series, "fields", {}).get(name, [])
    array = np.asarray(values, dtype=float)
    count = len(getattr(series, "times", ()))
    if array.shape != (count,):
        result = np.full(count, np.nan, dtype=float)
        flat = array.reshape(-1) if array.size else array
        result[: min(count, len(flat))] = flat[:count]
        return result
    return array


def _statistic(series: Any, parameter: str, statistic: str) -> np.ndarray:
    try:
        values = series.statistic(parameter, statistic)
    except Exception:
        values = getattr(series, "stats", {}).get(parameter, {}).get(statistic, [])
    array = np.asarray(values, dtype=float)
    count = len(getattr(series, "times", ()))
    if array.shape != (count,):
        result = np.full(count, np.nan, dtype=float)
        flat = array.reshape(-1) if array.size else array
        result[: min(count, len(flat))] = flat[:count]
        return result
    return array


def _cloud_values(series: Any) -> np.ndarray:
    total = _values(series, "cloud_cover")
    if np.isfinite(total).any():
        return total
    layers = np.vstack(
        [
            _values(series, "cloud_cover_low"),
            _values(series, "cloud_cover_mid"),
            _values(series, "cloud_cover_high"),
        ]
    )
    with np.errstate(all="ignore"):
        result = np.nanmax(layers, axis=0)
    result[np.all(~np.isfinite(layers), axis=0)] = np.nan
    return result


def _wind_direction_label(value: float | None) -> str:
    if value is None or not math.isfinite(value):
        return "направление —"
    index = int((float(value) % 360.0 + 11.25) // 22.5) % 16
    return WIND_DIRECTIONS_RU[index]


def _format_range(
    low: float | None,
    high: float | None,
    unit: str,
    *,
    signed: bool = False,
) -> str:
    if low is None or high is None:
        return "—"
    formatter = _fmt_signed if signed else _fmt
    if abs(high - low) < 0.05:
        return f"{formatter(low, 1)} {unit}"
    return f"{formatter(low, 1)}…{formatter(high, 1)} {unit}"


def _fmt(value: float | int | None, decimals: int = 1) -> str:
    if value is None:
        return "—"
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return "—"
    if not math.isfinite(numeric):
        return "—"
    text = f"{numeric:.{decimals}f}"
    if decimals == 0:
        text = str(int(round(numeric)))
    return text.replace("-", "−").replace(".", ",")


def _fmt_signed(value: float | int | None, decimals: int = 1) -> str:
    if value is None:
        return "—"
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return "—"
    if not math.isfinite(numeric):
        return "—"
    text = f"{numeric:+.{decimals}f}"
    return text.replace("-", "−").replace(".", ",")


def _finite_at(values: np.ndarray, index: int) -> float | None:
    if index < 0 or index >= len(values):
        return None
    value = float(values[index])
    return value if math.isfinite(value) else None


def _float_or_none(value: Any) -> float | None:
    try:
        result = float(value)
    except (TypeError, ValueError):
        return None
    return result if math.isfinite(result) else None


def _nanmin(values: np.ndarray) -> float | None:
    finite = np.asarray(values, dtype=float)
    finite = finite[np.isfinite(finite)]
    return float(np.min(finite)) if finite.size else None


def _nanmax(values: np.ndarray) -> float | None:
    finite = np.asarray(values, dtype=float)
    finite = finite[np.isfinite(finite)]
    return float(np.max(finite)) if finite.size else None


def _nanmean(values: np.ndarray) -> float | None:
    finite = np.asarray(values, dtype=float)
    finite = finite[np.isfinite(finite)]
    return float(np.mean(finite)) if finite.size else None


def _nansum(values: np.ndarray) -> float:
    finite = np.asarray(values, dtype=float)
    return float(np.nansum(finite)) if finite.size else 0.0


def _nanargmin(values: np.ndarray) -> int | None:
    array = np.asarray(values, dtype=float)
    if not np.isfinite(array).any():
        return None
    return int(np.nanargmin(array))


def _nanargmax(values: np.ndarray) -> int | None:
    array = np.asarray(values, dtype=float)
    if not np.isfinite(array).any():
        return None
    return int(np.nanargmax(array))


def _clean_text(value: Any) -> str:
    return " ".join(str(value or "").split())


def _safe_filename(value: str) -> str:
    text = _clean_text(value)
    text = re.sub(r"[^0-9A-Za-zА-Яа-яЁё._-]+", "_", text)
    text = re.sub(r"_+", "_", text).strip("._-")
    return (text or "точка")[:80]


def _find_libreoffice() -> str | None:
    configured = os.getenv("LIBREOFFICE_BIN", "").strip()
    if configured:
        path = shutil.which(configured) or (configured if Path(configured).is_file() else None)
        if path:
            return str(path)
    return shutil.which("soffice") or shutil.which("libreoffice")


def _rgb(value: str):
    from docx.shared import RGBColor

    return RGBColor.from_string(value)


def _set_style_font(style: Any, font_name: str) -> None:
    from docx.oxml.ns import qn

    style._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style._element.rPr.rFonts.set(qn("w:cs"), font_name)


def _set_cell_shading(cell: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell: Any, **kwargs: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin in ("top", "start", "bottom", "end"):
        if margin not in kwargs:
            continue
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(kwargs[margin]))
        node.set(qn("w:type"), "dxa")


def _configure_table(table: Any, headers: Sequence[str], widths_mm: Sequence[int]) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt

    table.autofit = False
    header_cells = table.rows[0].cells
    row_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    row_pr.append(tbl_header)
    for cell, text, width in zip(header_cells, headers, widths_mm):
        cell.width = Mm(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_shading(cell, "D9E8EF")
        _set_cell_margins(cell, top=70, start=70, bottom=70, end=70)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = REPORT_FONT
        run.font.size = Pt(7.1)
        run.font.color.rgb = _rgb("17324D")


def _fill_table_row(cells: Sequence[Any], values: Sequence[str], *, font_size: float) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    row = cells[0]._tc.getparent()
    tr_pr = row.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)
    for cell, text in zip(cells, values):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_margins(cell, top=55, start=65, bottom=55, end=65)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        lines = str(text).split("\n")
        for index, line in enumerate(lines):
            if index:
                paragraph.add_run().add_break()
            run = paragraph.add_run(line)
            run.font.name = REPORT_FONT
            run.font.size = Pt(font_size)


def _append_field(paragraph: Any, instruction: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    run.font.name = REPORT_FONT
    run.font.size = _pt(7)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))


def _pt(value: float):
    from docx.shared import Pt

    return Pt(value)
